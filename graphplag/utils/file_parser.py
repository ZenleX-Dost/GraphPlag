"""
File Parser Module for GraphPlag
Supports: PDF, DOCX, TXT, MD, and plain text
"""

import os
from pathlib import Path
from typing import Optional, Dict
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class FileParser:
    """Parse various file formats and extract text content"""
    
    SUPPORTED_FORMATS = ['.pdf', '.docx', '.txt', '.md', '.markdown']
    
    def __init__(self):
        """Initialize the file parser"""
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if required libraries are available"""
        self.has_pdf = False
        self.has_docx = False
        self.has_markdown = False
        
        try:
            import PyPDF2
            self.has_pdf = True
        except ImportError:
            logger.warning("PyPDF2 not installed. PDF support disabled.")
        
        try:
            import docx
            self.has_docx = True
        except ImportError:
            logger.warning("python-docx not installed. DOCX support disabled.")
        
        try:
            import markdown
            self.has_markdown = True
        except ImportError:
            logger.warning("markdown not installed. MD support disabled.")
    
    def parse_file(self, file_path: str) -> str:
        """
        Parse a file and extract text content
        
        Args:
            file_path: Path to the file to parse
            
        Returns:
            Extracted text content
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format is not supported
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get file extension
        ext = path.suffix.lower()
        
        if ext not in self.SUPPORTED_FORMATS:
            raise ValueError(
                f"Unsupported file format: {ext}. "
                f"Supported formats: {', '.join(self.SUPPORTED_FORMATS)}"
            )
        
        # Parse based on extension
        if ext == '.pdf':
            return self._parse_pdf(path)
        elif ext == '.docx':
            return self._parse_docx(path)
        elif ext in ['.txt', '.md', '.markdown']:
            return self._parse_text(path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    def _parse_pdf(self, path: Path) -> str:
        """Parse PDF file"""
        if not self.has_pdf:
            raise ImportError("PyPDF2 is required for PDF support. Install with: pip install PyPDF2")
        
        try:
            import PyPDF2
            
            text_content = []
            
            with open(path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract text from all pages
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text)
                
            full_text = '\n\n'.join(text_content)
            
            if not full_text.strip():
                logger.warning(f"No text extracted from PDF: {path}")
                return ""
            
            logger.info(f"Extracted {len(full_text)} characters from PDF: {path.name}")
            return full_text
            
        except Exception as e:
            logger.error(f"Error parsing PDF {path}: {e}")
            raise ValueError(f"Failed to parse PDF file: {e}")
    
    def _parse_docx(self, path: Path) -> str:
        """Parse DOCX file"""
        if not self.has_docx:
            raise ImportError("python-docx is required for DOCX support. Install with: pip install python-docx")
        
        try:
            import docx
            
            doc = docx.Document(path)
            
            # Extract text from paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            full_text = '\n\n'.join(text_content)
            
            if not full_text.strip():
                logger.warning(f"No text extracted from DOCX: {path}")
                return ""
            
            logger.info(f"Extracted {len(full_text)} characters from DOCX: {path.name}")
            return full_text
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {path}: {e}")
            raise ValueError(f"Failed to parse DOCX file: {e}")
    
    def _parse_text(self, path: Path) -> str:
        """Parse text file (TXT, MD)"""
        try:
            # Try UTF-8 first
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(path, 'r', encoding=encoding) as file:
                        text = file.read()
                    
                    if text.strip():
                        logger.info(f"Extracted {len(text)} characters from {path.name} (encoding: {encoding})")
                        return text
                    
                except UnicodeDecodeError:
                    continue
            
            raise ValueError(f"Could not decode file with any supported encoding")
            
        except Exception as e:
            logger.error(f"Error parsing text file {path}: {e}")
            raise ValueError(f"Failed to parse text file: {e}")
    
    def parse_from_bytes(self, file_bytes: bytes, filename: str) -> str:
        """
        Parse file from bytes (useful for uploaded files)
        
        Args:
            file_bytes: File content as bytes
            filename: Original filename with extension
            
        Returns:
            Extracted text content
        """
        import tempfile
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as tmp_file:
            tmp_file.write(file_bytes)
            tmp_path = tmp_file.name
        
        try:
            # Parse the temporary file
            text = self.parse_file(tmp_path)
            return text
        finally:
            # Clean up temporary file
            try:
                os.unlink(tmp_path)
            except:
                pass
    
    def get_file_info(self, file_path: str) -> Dict[str, any]:
        """
        Get information about a file
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary with file information
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        return {
            'name': path.name,
            'extension': path.suffix.lower(),
            'size_bytes': path.stat().st_size,
            'size_mb': path.stat().st_size / (1024 * 1024),
            'is_supported': path.suffix.lower() in self.SUPPORTED_FORMATS,
        }
    
    @staticmethod
    def is_supported(filename: str) -> bool:
        """Check if a file format is supported"""
        ext = Path(filename).suffix.lower()
        return ext in FileParser.SUPPORTED_FORMATS


# Convenience function
def parse_document(file_path: str) -> str:
    """
    Quick function to parse a document
    
    Args:
        file_path: Path to the document
        
    Returns:
        Extracted text content
    """
    parser = FileParser()
    return parser.parse_file(file_path)


# Example usage
if __name__ == "__main__":
    parser = FileParser()
    
    # Test with a sample file
    test_file = "sample.txt"
    if os.path.exists(test_file):
        text = parser.parse_file(test_file)
        print(f"Extracted {len(text)} characters")
        print(f"Preview: {text[:200]}...")
